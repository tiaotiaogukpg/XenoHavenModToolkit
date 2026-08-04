# -*- coding: utf-8 -*-
"""导出《异星家园 Mod 简易说明》为 Word，并按配图文件夹插入截图或占位标记。"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
IMAGE_DIR = ROOT / "docs" / "简易说明配图"
OUTPUT = ROOT / "docs" / "异星家园Mod简易说明.docx"

# 文件名 → 简短说明（用于占位框与图注）
FIGURES = {
    "01-main-window.png": "工具完整主窗口（顶栏 / 左侧树 / 右侧磁贴 / 底部日志）",
    "02-toolbar-buttons.png": "顶栏工具按钮一览",
    "03-new-mod-dialog.png": "「新建MOD」对话框",
    "04-building-editor.png": "组件（Building）编辑窗口",
    "05-steam-connected.png": "顶栏 Steam 已连接状态",
    "06-workshop-upload.png": "「上传到创意工坊」窗口",
    "07-workshop-subscribe.png": "Steam 工坊页面「订阅」按钮",
    "08-game-mod-list.png": "游戏内 Mod 列表 / 启用界面",
}


def set_run_font(run, *, size_pt: float = 11, bold: bool = False, color: RGBColor | None = None) -> None:
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size_pt=16 if level == 1 else 13, bold=True)


def add_para(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run)
    p.paragraph_format.space_after = Pt(3)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    set_run_font(run)
    p.paragraph_format.space_after = Pt(3)


def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, bold=True, size_pt=10)
        set_cell_shading(cell, "D9E2F3")
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(value)
            set_run_font(run, size_pt=10)
    doc.add_paragraph()


def add_figure(doc: Document, filename: str, caption_no: int) -> None:
    desc = FIGURES[filename]
    path = IMAGE_DIR / filename
    caption = f"图 {caption_no}　{desc}"

    if path.is_file() and path.stat().st_size > 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Cm(14.5))
    else:
        # 占位框：醒目提示补图位置与文件名
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        cell = table.rows[0].cells[0]
        set_cell_shading(cell, "FFF2CC")
        cell.text = ""
        lines = [
            "【此处插入操作截图】",
            f"请将图片放到：docs/简易说明配图/{filename}",
            f"建议内容：{desc}",
        ]
        for i, line in enumerate(lines):
            para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(line)
            set_run_font(
                run,
                size_pt=11 if i == 0 else 10,
                bold=(i == 0),
                color=RGBColor(0xC0, 0x00, 0x00) if i == 0 else RGBColor(0x59, 0x59, 0x59),
            )

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run_font(run, size_pt=9, color=RGBColor(0x59, 0x59, 0x59))
    cap.paragraph_format.space_after = Pt(12)


def build() -> Path:
    IMAGE_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("异星家园（XenoHaven）Mod 简易说明")
    set_run_font(run, size_pt=20, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("玩家向 · 可做类型 / 界面功能 / 字段说明 / 上传与下载")
    set_run_font(run, size_pt=11, color=RGBColor(0x59, 0x59, 0x59))

    add_para(
        doc,
        "《异星家园》支持玩家自制 Mod，并通过 Steam 创意工坊分享与订阅。"
        "制作请使用配套工具 XenoHaven MOD Toolkit（XenoHavenModTool.exe）。",
    )
    add_para(doc, "创意工坊入口：https://steamcommunity.com/workshop/about/?appid=3461270")

    note = doc.add_paragraph()
    run = note.add_run(
        "配图说明：正文中黄色方框为截图占位。请按文件名把图片补到 docs/简易说明配图/ 后，"
        "再运行 scripts/export-mod-guide-docx.py 重新导出本文档。"
    )
    set_run_font(run, size_pt=9, color=RGBColor(0xC0, 0x00, 0x00))
    note.paragraph_format.space_after = Pt(12)

    # —— 1 ——
    add_heading(doc, "1. 可以制作哪些 Mod", level=1)
    add_para(doc, "当前版本主要支持建筑类 Mod（Thing/Buildings），可制作的组件类型包括：")
    add_table(
        doc,
        ["类型", "说明"],
        [
            ["BOX", "储物箱等有容量的容器"],
            ["SIMPLE_OBJECT", "装饰物、摆件"],
            ["SMALL_LAMP", "小型灯具"],
            ["STREET_LIGHT", "路灯"],
            ["PRODUCTION_LINE", "自定义外观的生产线（需选择模拟的原版生产线）"],
        ],
    )
    add_para(doc, "暂不支持脚本、Prefab 编辑等更复杂扩展。")

    # —— 2 ——
    add_heading(doc, "2. 工具界面功能介绍", level=1)
    add_para(doc, "启动工具后，主界面大致分为四块：")
    add_bullet(doc, "顶栏：常用操作按钮")
    add_bullet(doc, "左侧树：Mods 目录下的工程列表与文件结构")
    add_bullet(doc, "右侧：当前 Mod 的组件磁贴预览")
    add_bullet(doc, "底部：日志 / 校验结果")

    add_figure(doc, "01-main-window.png", 1)

    add_heading(doc, "2.1 顶栏主要功能", level=2)
    add_figure(doc, "02-toolbar-buttons.png", 2)
    add_table(
        doc,
        ["按钮", "作用"],
        [
            ["打开工程", "打开任意包含 main.xml 的 Mod 文件夹"],
            ["新建MOD", "创建新的建筑类 Mod（填写名称、作者、描述，并导入图标与截图）"],
            ["删除MOD", "删除当前选中的 Mod 工程（不可恢复）"],
            ["编辑MOD", "修改 Mod 名称、作者、版本、描述、图标、截图等"],
            ["上传工坊", "将当前 Mod 发布或更新到 Steam 创意工坊（需已登录 Steam）"],
            ["新建组件", "在当前 Mod 中新增一个建筑组件"],
            ["编辑组件", "修改选中组件的属性、材料、图片等"],
            ["删除组件", "删除选中的组件"],
            ["重连 Steam", "重新连接本机 Steam 客户端"],
        ],
    )
    add_para(
        doc,
        "左侧点选不同节点时，顶栏按钮会相应启用或变灰（类似资源管理器）。"
        "右侧磁贴可单击选中、双击编辑组件。",
    )

    add_heading(doc, "2.2 相关操作窗口示意", level=2)
    add_para(doc, "新建 Mod 时会弹出信息填写窗口：")
    add_figure(doc, "03-new-mod-dialog.png", 3)
    add_para(doc, "新建或编辑组件时会打开 Building 编辑窗口：")
    add_figure(doc, "04-building-editor.png", 4)

    # —— 3 字段说明 ——
    add_heading(doc, "3. Mod / Building 字段说明", level=1)
    add_para(
        doc,
        "新建或编辑时会看到两类字段：Mod 信息（整个 Mod 的身份证）"
        "和 Building 组件（单个建筑的属性）。",
    )

    add_heading(doc, "3.1 Mod 字段（新建 / 编辑 MOD）", level=2)
    add_table(
        doc,
        ["字段", "说明"],
        [
            ["id", "Mod 基础 ID，创建时自动生成，之后不可修改。游戏内最终物品 ID ≈ 该值 + 组件本地序号。"],
            ["steamPublishedFileId", "Steam 创意工坊作品 ID。新建时为 0；首次上传成功后由工具自动写入，用于以后更新同一条作品。"],
            ["SupportVersion", "支持的游戏版本标记，当前固定为 1。"],
            ["name", "Mod 显示名称（必填），会出现在游戏列表和工坊标题中。"],
            ["auth", "作者名（必填）。"],
            ["version", "Mod 版本号（必填），默认 1.0.0，内容更新时可自行递增。"],
            ["Category", "分类，默认 Building，用于工具总览分组。"],
            ["description", "Mod 描述（必填），工坊简介也会用到。"],
            ["icon.png", "Mod 列表图标（必填）。"],
            ["screenshot.png", "Mod 详情 / 工坊预览图（必填）。"],
        ],
    )

    add_heading(doc, "3.2 Building 字段（组件编辑窗口）", level=2)
    add_para(doc, "工具里的「组件」= 一条 Building 定义。主要字段如下：")
    add_table(
        doc,
        ["字段", "说明"],
        [
            ["id", "组件本地序号（1、2、3…），只读；同时决定图片文件名。不要改成「Mod基础ID+序号」。"],
            ["name", "组件显示名称（必填）。"],
            ["type", "建筑类型：BOX / SIMPLE_OBJECT / SMALL_LAMP / STREET_LIGHT / PRODUCTION_LINE。"],
            ["direction", "朝向。箱子、装饰物、灯类固定为 1；生产线可选 1 或 3。"],
            ["workbenchId", "在哪张工作台制造。请从下拉列表选择（界面显示「名称-ID」）。"],
            ["simulateId", "仅 PRODUCTION_LINE 需要：选择要模拟的原版生产线。"],
            ["capbility", "容量，范围 16～96。装饰物 / 灯 / 生产线不显示此字段。拼写固定为 capbility（不要写成 capability）。"],
            ["health", "固定为 10，一般无需修改。"],
            ["size.x / size.y", "占地格子数，须为正整数。"],
            ["制造公式（materials）", "建造消耗的材料列表；材料从下拉选择，单条数量 1～200。"],
            ["碰撞（barrier）", "开启后角色无法穿过；BOX 默认开启，其它类型默认关闭。"],
            ["组件图片", "地图上显示的外观，保存为 images/<id>.png。"],
            ["物品栏图标", "背包中的小图标，保存为 images/icon/<id>.png。"],
        ],
    )
    add_para(doc, "类型与字段关系简表：")
    add_table(
        doc,
        ["type", "direction", "容量", "simulateId", "默认碰撞"],
        [
            ["BOX", "固定 1", "需要（16～96）", "无", "开"],
            ["SIMPLE_OBJECT", "固定 1", "无", "无", "关"],
            ["SMALL_LAMP", "固定 1", "无", "无", "关"],
            ["STREET_LIGHT", "固定 1", "无", "无", "关"],
            ["PRODUCTION_LINE", "1 或 3", "无", "必填", "编辑时默认开"],
        ],
    )

    # —— 4 ——
    add_heading(doc, "4. 上传与下载", level=1)

    add_heading(doc, "4.1 上传（发布 / 更新自己的 Mod）", level=2)
    add_numbered(doc, "启动并登录 Steam 客户端，确认工具顶栏显示已连接 Steam。")
    add_figure(doc, "05-steam-connected.png", 5)
    add_numbered(doc, "在工具中打开要发布的 Mod。")
    add_numbered(doc, "确认已填写名称、描述，并放好 icon.png、screenshot.png。")
    add_numbered(doc, "点击「上传工坊」，填写标题、简介，选择可见性（公开 / 仅好友 / 私有）。")
    add_figure(doc, "06-workshop-upload.png", 6)
    add_numbered(doc, "建议勾选「上传时排除 .meta 文件」，然后开始上传。")

    add_para(doc, "说明：", bold=True)
    add_bullet(doc, "首次上传会创建新的工坊条目；成功后工具会把工坊 ID 写回 Mod。")
    add_bullet(doc, "之后再点「上传工坊」即更新同一条作品，无需重新新建。")
    add_bullet(doc, "若 Steam 提示需接受创意工坊协议，按页面提示同意后再公开。")

    add_heading(doc, "4.2 下载（使用他人的 Mod）", level=2)
    add_numbered(
        doc,
        "打开游戏的 Steam 创意工坊页面：https://steamcommunity.com/workshop/about/?appid=3461270",
    )
    add_numbered(doc, "浏览或搜索喜欢的作品，点击「订阅」。")
    add_figure(doc, "07-workshop-subscribe.png", 7)
    add_numbered(doc, "Steam 会自动下载该 Mod。")
    add_numbered(doc, "启动《异星家园》，在游戏的 Mod 列表中启用对应 Mod 即可。")
    add_figure(doc, "08-game-mod-list.png", 8)
    add_para(doc, "取消订阅后，Steam 一般会移除对应内容；具体以游戏内 Mod 管理界面为准。")

    # 附录：配图清单
    add_heading(doc, "附录：配图文件清单", level=1)
    add_para(doc, "请将下列文件放到文件夹：docs/简易说明配图/")
    add_table(
        doc,
        ["文件名", "建议截图内容"],
        [[name, desc] for name, desc in FIGURES.items()],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    return OUTPUT


if __name__ == "__main__":
    out = build()
    missing = [name for name in FIGURES if not (IMAGE_DIR / name).is_file()]
    print(f"Wrote: {out}")
    if missing:
        print(f"Missing images ({len(missing)}), placeholders kept:")
        for name in missing:
            print(f"  - {name}")
    else:
        print("All images present.")
